#!/usr/bin/env python3
"""MicroHydroV1 - Auto-fill RUN-3 Evidence Doc (v0.3.0+)

Adds to previous version:
- Pulls T-001, T-003, T-004 latest value + pass/fail from automation/MicroHydroV1_RnD_Export.xlsx
- Populates those cells in the Evidence table automatically.

Inputs
- T-002 timeseries CSV (Time_s, Pressure_Pa)
- RnD workbook (default: automation/MicroHydroV1_RnD_Export.xlsx)

Outputs
- Populated evidence DOCX
- Optional PSD plot PNG (delegates to tools/plots/plot_t002_psd_peakiness.py)

Usage
  python tools/workflow/autofill_evidence_doc.py     --root .     --run-id 2026-01-22_RUN3     --t002 tests/raw/RUN3/T002_TankRipple_timeseries.csv     --out-doc docs/run3/RUN3_Evidence_2026-01-22_RUN3.docx     --out-plot tests/results/summary/T002_PSD_Peakiness.png

Optional
  --workbook automation/MicroHydroV1_RnD_Export.xlsx

Notes
- openpyxl does not calculate formulas; if Pass cells are blank, this script evaluates PASS/FAIL
  using the project acceptance rules for T-001/T-003/T-004.
"""

from __future__ import annotations

import argparse
import datetime as dt
from pathlib import Path
from typing import Dict, Any, Optional

import numpy as np
import pandas as pd
from docx import Document

BAND_LOW = 0.2
BAND_HIGH = 20.0


def compute_psd(x: np.ndarray, fs: float):
    x = x - np.nanmean(x)
    try:
        from scipy.signal import welch
        f, pxx = welch(x, fs=fs, nperseg=min(2048, len(x)))
    except Exception:
        n = len(x)
        win = np.hanning(n)
        xf = np.fft.rfft(x * win)
        pxx = (np.abs(xf) ** 2) / (fs * np.sum(win ** 2))
        f = np.fft.rfftfreq(n, d=1.0 / fs)
    return f, pxx


def peakiness_metrics(csv_path: Path, time_col='Time_s', signal_col='Pressure_Pa'):
    df = pd.read_csv(csv_path)
    if time_col not in df.columns or signal_col not in df.columns:
        raise SystemExit(f"Missing columns. Need {time_col} and {signal_col}. Got: {list(df.columns)}")

    t = pd.to_numeric(df[time_col], errors='coerce').to_numpy()
    x = pd.to_numeric(df[signal_col], errors='coerce').to_numpy()
    m = np.isfinite(t) & np.isfinite(x)
    t, x = t[m], x[m]
    if len(t) < 10:
        raise SystemExit('Not enough samples after cleaning.')

    dtv = np.diff(t)
    dtv = dtv[np.isfinite(dtv) & (dtv > 0)]
    fs = 1.0 / float(np.median(dtv)) if len(dtv) else 100.0

    f, pxx = compute_psd(x, fs)
    band = (f >= BAND_LOW) & (f <= BAND_HIGH)
    if not np.any(band):
        raise SystemExit('No PSD bins in band. Check sampling/timebase.')

    p = pxx[band]
    peak = float(np.max(p))
    med = float(np.median(p))
    peakiness = peak / max(1e-12, med)
    f_dom = float(f[band][np.argmax(p)])

    return {
        'fs': fs,
        'peakiness': peakiness,
        'f_dom': f_dom,
        'peak_psd': peak,
        'median_psd': med,
    }


def _norm_test_id(s: str) -> str:
    if s is None:
        return ''
    # normalize unicode hyphens to '-'
    return str(s).replace('‑', '-').replace('–', '-').replace('−', '-').strip()


def _safe_float(x) -> Optional[float]:
    try:
        if x is None or x == '':
            return None
        return float(x)
    except Exception:
        return None


def _parse_ts(x) -> Optional[dt.datetime]:
    if x is None or x == '':
        return None
    if isinstance(x, dt.datetime):
        return x
    s = str(x).strip()
    # try ISO first
    try:
        return dt.datetime.fromisoformat(s)
    except Exception:
        pass
    # try pandas parser
    try:
        return pd.to_datetime(s).to_pydatetime()
    except Exception:
        return None


def eval_pass_rule(test_id: str, sensor: str, value, unit: str) -> Optional[bool]:
    tid = _norm_test_id(test_id)
    s = (sensor or '').lower()
    v = _safe_float(value)

    if v is None:
        return None

    if tid == 'T-001':
        # coherence >= 0.90
        if 'coherence' in s or s == 'coherence':
            return v >= 0.90
        return v >= 0.90

    if tid == 'T-003':
        # drift <= 0.1 Hz
        if 'elc' in s or 'drift' in s or 'stability' in s:
            return v <= 0.10
        return v <= 0.10

    if tid == 'T-004':
        # power 10-12 kW
        if 'power' in s:
            return 10.0 <= v <= 12.0
        return 10.0 <= v <= 12.0

    return None


def read_workbook_latest(workbook_path: Path) -> Dict[str, Dict[str, Any]]:
    """Read Measurements sheet and return latest row per Test_ID.

    Returns:
      {'T-001': {'value':..., 'unit':..., 'pass':..., 'timestamp':...}, ...}
    """
    try:
        import openpyxl
    except Exception as e:
        raise SystemExit('openpyxl required to read workbook')

    wb = openpyxl.load_workbook(workbook_path, data_only=True)
    if 'Measurements' not in wb.sheetnames:
        raise SystemExit('Workbook missing Measurements sheet: ' + str(workbook_path))
    ws = wb['Measurements']

    # map headers
    headers = {}
    for c in range(1, ws.max_column + 1):
        v = ws.cell(1, c).value
        if v:
            headers[str(v).strip()] = c

    def col(name, default=None):
        return headers.get(name, default)

    idx_ts = col('Timestamp')
    idx_tid = col('Test_ID')
    idx_sensor = col('Sensor')
    idx_val = col('Value')
    idx_unit = col('Unit')
    idx_pass = col('Pass')

    if not all([idx_ts, idx_tid, idx_sensor, idx_val]):
        raise SystemExit('Measurements sheet missing required headers.')

    best: Dict[str, Dict[str, Any]] = {}

    for r in range(2, ws.max_row + 1):
        tid = ws.cell(r, idx_tid).value
        if not tid:
            continue
        tid_n = _norm_test_id(tid)
        if tid_n not in ('T-001','T-003','T-004'):
            continue

        ts = _parse_ts(ws.cell(r, idx_ts).value)
        sensor = ws.cell(r, idx_sensor).value
        val = ws.cell(r, idx_val).value
        unit = ws.cell(r, idx_unit).value if idx_unit else ''
        pas = ws.cell(r, idx_pass).value if idx_pass else None

        # normalize pass
        if isinstance(pas, str):
            pnorm = pas.strip().lower()
            if pnorm in ('true','pass','yes','1'):
                pas = True
            elif pnorm in ('false','fail','no','0'):
                pas = False
            else:
                pas = None

        if pas is None:
            pas = eval_pass_rule(tid_n, str(sensor or ''), val, str(unit or ''))

        record = {
            'timestamp': ts,
            'sensor': sensor,
            'value': val,
            'unit': unit,
            'pass': pas,
        }

        if tid_n not in best:
            best[tid_n] = record
        else:
            # pick latest timestamp; if missing timestamps, keep first
            old_ts = best[tid_n].get('timestamp')
            if ts and (old_ts is None or ts > old_ts):
                best[tid_n] = record

    return best


def replace_in_paragraph(paragraph, mapping: dict[str, str]):
    # simple full-text replace on paragraph runs
    for key, val in mapping.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(key, val)


def fill_table_values(doc: Document, metrics: Dict[str, Dict[str, Any]], peakiness: str, f_dom: str):
    """Populate the summary table for T-001/T-003/T-004 and resonance row."""

    for table in doc.tables:
        # identify header row
        if len(table.rows) < 2:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        if 'Test' not in header or 'Result' not in header:
            continue

        # columns
        col_test = 0
        col_metric = header.index('Metric') if 'Metric' in header else 1
        col_result = header.index('Result') if 'Result' in header else 3
        col_pass = header.index('Pass/Fail') if 'Pass/Fail' in header else 4

        for row in table.rows[1:]:
            test_cell = _norm_test_id(row.cells[col_test].text)
            metric_cell = row.cells[col_metric].text.strip()

            # resonance row
            if test_cell == 'T-002' and 'Resonance' in metric_cell:
                row.cells[col_result].text = peakiness
                row.cells[col_pass].text = 'n/a'
                continue

            if test_cell in metrics:
                rec = metrics[test_cell]
                v = _safe_float(rec.get('value'))
                unit = str(rec.get('unit') or '').strip()
                pas = rec.get('pass')

                if v is not None:
                    if test_cell == 'T-001':
                        val_str = f"{v:.3f}"
                    elif test_cell == 'T-003':
                        val_str = f"{v:.3f}"
                    else:
                        val_str = f"{v:.2f}"
                else:
                    val_str = str(rec.get('value') or '').strip() or '____'

                if unit:
                    val_str = f"{val_str} {unit}"

                row.cells[col_result].text = val_str

                if pas is True:
                    row.cells[col_pass].text = 'PASS'
                elif pas is False:
                    row.cells[col_pass].text = 'FAIL'
                else:
                    row.cells[col_pass].text = 'TBD'

    # also append a short summary paragraph
    try:
        doc.add_paragraph('')
        p = doc.add_paragraph('Auto-filled from workbook:')
        for tid in ('T-001','T-003','T-004'):
            if tid in metrics:
                rec = metrics[tid]
                ts = rec.get('timestamp')
                ts_s = ts.isoformat(timespec='seconds') if ts else ''
                v = rec.get('value')
                unit = rec.get('unit') or ''
                pas = rec.get('pass')
                p2 = doc.add_paragraph(f"- {tid}: {v} {unit} | {('PASS' if pas else 'FAIL' if pas is False else 'TBD')} | {ts_s}")
        doc.add_paragraph(f"- T-002 peakiness: {peakiness} | dominant Hz: {f_dom}")
    except Exception:
        pass


def main(argv=None):
    ap = argparse.ArgumentParser()
    ap.add_argument('--root', default='.', help='MicroHydroV1 root')
    ap.add_argument('--run-id', required=True)
    ap.add_argument('--t002', required=True, help='Path to T-002 timeseries CSV')
    ap.add_argument('--out-doc', required=True, help='Output populated docx')
    ap.add_argument('--template', default='docs/run3/RUN3_Evidence_Template.docx')
    ap.add_argument('--out-plot', default=None, help='Optional: output plot PNG by calling plot script')
    ap.add_argument('--workbook', default='automation/MicroHydroV1_RnD_Export.xlsx', help='RnD workbook path')
    args = ap.parse_args(argv)

    root = Path(args.root).resolve()

    # T-002 metrics
    t002_path = Path(args.t002)
    if not t002_path.is_absolute():
        t002_path = root / t002_path

    t2 = peakiness_metrics(t002_path)
    peakiness = f"{t2['peakiness']:.3f}"
    f_dom = f"{t2['f_dom']:.2f}"
    fs = f"{t2['fs']:.2f}"

    # optional plot
    if args.out_plot:
        plot_script = root/'tools/plots/plot_t002_psd_peakiness.py'
        if plot_script.exists():
            import subprocess
            out_plot = Path(args.out_plot)
            if not out_plot.is_absolute():
                out_plot = root/out_plot
            out_plot.parent.mkdir(parents=True, exist_ok=True)
            subprocess.check_call(['python', str(plot_script), '--csv', str(t002_path), '--out', str(out_plot)])

    # workbook summary for T-001/T-003/T-004
    wb_path = Path(args.workbook)
    if not wb_path.is_absolute():
        wb_path = root / wb_path
    metrics = {}
    if wb_path.exists():
        metrics = read_workbook_latest(wb_path)

    # doc template
    template = root / args.template
    if not template.exists():
        raise SystemExit('Template not found: ' + str(template))

    doc = Document(str(template))
    now = dt.datetime.now().isoformat(timespec='seconds')

    mapping = {
        '{{RUN_ID}}': args.run_id,
        '{{DATETIME}}': now,
        '{{PEAKINESS}}': peakiness,
        '{{F_DOM}}': f_dom,
        'Observed dominant peak (Hz): ____': f"Observed dominant peak (Hz): {f_dom}",
        'Sampling: 100 Hz': f"Sampling: 100 Hz (inferred fs≈{fs} Hz)",
    }

    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)

    fill_table_values(doc, metrics=metrics, peakiness=peakiness, f_dom=f_dom)

    out_doc = Path(args.out_doc)
    if not out_doc.is_absolute():
        out_doc = root/out_doc
    out_doc.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_doc))

    print('Wrote:', out_doc)
    print('T-002 peakiness:', peakiness, 'dominant Hz:', f_dom, 'fs≈', fs)
    for k,v in metrics.items():
        print('Workbook', k, 'value=', v.get('value'), 'pass=', v.get('pass'))


if __name__ == '__main__':
    main()
