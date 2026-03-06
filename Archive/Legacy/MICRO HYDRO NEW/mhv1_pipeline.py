#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import argparse, sys, os, json, subprocess
from pathlib import Path
from datetime import datetime
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document

ACCEPT = {
  'T001': {'min': 0.90},
  'T002': {'ratio_low': 0.50, 'ratio_high': 0.70},
  'T003': {'max_drift': 0.10},
  'T004': {'p_lo': 10.0, 'p_hi': 12.0},
}

PRIO_WORKBOOKS = [
  'MicroHydroV1_RnD_FullUpgrade.xlsx',
  'MicroHydroV1_RnD_Export.xlsx',
  'MicroHydroV1_RnD_Export (1).xlsx',
]

PLOT_DIR_REL = '03_data/analysis/plots'

def find_first(root: Path, names):
    for nm in names:
        for p in root.rglob(nm):
            return p
    return None

def find_rnd_workbook(root: Path):
    p = find_first(root, PRIO_WORKBOOKS)
    if p:
        return p
    for p in root.rglob('*.xlsx'):
        if 'RnD' in p.name or 'Tracker' in p.name:
            return p
    return None

def discover_intake_csvs(root: Path):
    csvs = []
    for d in root.rglob('intake_raw'):
        if d.is_dir():
            csvs.extend(sorted(d.glob('*.csv')))
    return sorted(csvs)

def import_csvs_to_workbook(csvs, workbook: Path):
    count = 0
    try:
        xl = pd.ExcelFile(workbook, engine='openpyxl')
        existing = pd.read_excel(workbook, sheet_name='Measurements', engine='openpyxl') if 'Measurements' in xl.sheet_names else None
    except Exception:
        existing = None
    combined = existing if existing is not None else pd.DataFrame()
    for c in csvs:
        try:
            df = pd.read_csv(c)
            needed = ['Timestamp','Test_ID','Sensor','Value','Unit','Baseline_Value','Conditions','Notes']
            for col in needed:
                if col not in df.columns:
                    df[col] = ''
            combined = pd.concat([combined, df], ignore_index=True)
            count += len(df)
        except Exception as e:
            print(f"[WARN] Skipping CSV {c}: {e}")
    with pd.ExcelWriter(workbook, engine='openpyxl', mode='a', if_sheet_exists='replace') as w:
        combined.to_excel(w, sheet_name='Measurements', index=False)
    return count

def compute_metrics(df: pd.DataFrame):
    out = {}
    df = df.copy()
    df['Sensor'] = df['Sensor'].astype(str)
    df['Test_ID'] = df['Test_ID'].astype(str)
    # T001 coherence
    t1 = df[df['Test_ID'].str.contains('T-?001', case=False) & df['Sensor'].str.contains('coherence', case=False)]
    if not t1.empty:
        vals = pd.to_numeric(t1['Value'], errors='coerce').dropna()
        out['T001'] = {
            'count': int(vals.size),
            'min': float(vals.min()) if not vals.empty else None,
            'max': float(vals.max()) if not vals.empty else None,
            'mean': float(vals.mean()) if not vals.empty else None,
            'pass_rate': float((vals >= ACCEPT['T001']['min']).mean()) if not vals.empty else 0.0,
        }
    # T002 ripple ratio
    t2 = df[df['Test_ID'].str.contains('T-?002', case=False) & df['Sensor'].str.contains('ripple', case=False)].copy()
    if not t2.empty:
        t2['Value'] = pd.to_numeric(t2['Value'], errors='coerce')
        t2['Baseline_Value'] = pd.to_numeric(t2['Baseline_Value'], errors='coerce')
        t2 = t2.dropna(subset=['Value','Baseline_Value'])
        t2['ratio'] = t2['Value'] / t2['Baseline_Value']
        out['T002'] = {
            'count': int(t2.shape[0]),
            'mean_mm': float(t2['Value'].mean()) if not t2.empty else None,
            'mean_baseline_mm': float(t2['Baseline_Value'].mean()) if not t2.empty else None,
            'mean_ratio': float(t2['ratio'].mean()) if not t2.empty else None,
            'pass_rate': float(((t2['ratio'] >= ACCEPT['T002']['ratio_low']) & (t2['ratio'] <= ACCEPT['T002']['ratio_high'])).mean()) if not t2.empty else 0.0,
        }
    # T003 drift
    t3 = df[df['Test_ID'].str.contains('T-?003', case=False) & df['Sensor'].str_contains('elc', case=False)]
    if not t3.empty:
        vals = pd.to_numeric(t3['Value'], errors='coerce').dropna()
        out['T003'] = {
            'count': int(vals.size),
            'min_Hz': float(vals.min()) if not vals.empty else None,
            'max_Hz': float(vals.max()) if not vals.empty else None,
            'mean_Hz': float(vals.mean()) if not vals.empty else None,
            'pass_rate': float((vals <= ACCEPT['T003']['max_drift']).mean()) if not vals.empty else 0.0,
        }
    # T004 power window
    t4 = df[df['Test_ID'].str.contains('T-?004', case=False) & df['Sensor'].str.contains('power', case=False)]
    if not t4.empty:
        vals = pd.to_numeric(t4['Value'], errors='coerce').dropna()
        out['T004'] = {
            'count': int(vals.size),
            'min_kW': float(vals.min()) if not vals.empty else None,
            'max_kW': float(vals.max()) if not vals.empty else None,
            'mean_kW': float(vals.mean()) if not vals.empty else None,
            'pass_rate': float(((vals >= ACCEPT['T004']['p_lo']) & (vals <= ACCEPT['T004']['p_hi'])).mean()) if not vals.empty else 0.0,
        }
    return out

def make_plots(df: pd.DataFrame, outdir: Path):
    outdir.mkdir(parents=True, exist_ok=True)
    plots = {}
    # T003
    t3 = df[df['Test_ID'].str.contains('T-?003', case=False) & df['Sensor'].str.contains('elc', case=False)].copy()
    if not t3.empty:
        t3['Timestamp'] = pd.to_datetime(t3['Timestamp'], errors='coerce')
        t3 = t3.dropna(subset=['Timestamp']).sort_values('Timestamp')
        plt.figure(figsize=(7,4))
        plt.plot(t3['Timestamp'], pd.to_numeric(t3['Value'], errors='coerce'), marker='o')
        plt.axhline(ACCEPT['T003']['max_drift'], color='r', linestyle='--', label='0.1 Hz')
        plt.title('T-003 ELC Drift vs Time (Hz)')
        plt.xlabel('Time'); plt.ylabel('Drift (Hz)'); plt.grid(True); plt.legend()
        p = outdir / 'T003_ELC_Drift_Time.png'
        plt.tight_layout(); plt.savefig(p, dpi=150); plt.close()
        plots['T003'] = str(p)
    # T004
    t4 = df[df['Test_ID'].str.contains('T-?004', case=False) & df['Sensor'].str.contains('power', case=False)].copy()
    if not t4.empty:
        t4['Timestamp'] = pd.to_datetime(t4['Timestamp'], errors='coerce')
        t4 = t4.dropna(subset=['Timestamp']).sort_values('Timestamp')
        vals = pd.to_numeric(t4['Value'], errors='coerce')
        plt.figure(figsize=(7,4))
        plt.plot(t4['Timestamp'], vals, marker='o')
        plt.axhline(ACCEPT['T004']['p_lo'], color='g', linestyle='--', label='10 kW')
        plt.axhline(ACCEPT['T004']['p_hi'], color='g', linestyle='--', label='12 kW')
        plt.title('T-004 Power vs Time (kW)')
        plt.xlabel('Time'); plt.ylabel('Power (kW)'); plt.grid(True); plt.legend()
        p = outdir / 'T004_Power_Time.png'
        plt.tight_layout(); plt.savefig(p, dpi=150); plt.close()
        plots['T004'] = str(p)
    return plots

def update_gonogo(docx_path: Path, metrics, plots):
    if not docx_path or not docx_path.exists():
        print(f"[WARN] Go/No-Go not found: {docx_path}")
        return
    doc = Document(docx_path)
    target = None
    for tbl in doc.tables:
        hdr = [c.text.strip() for c in tbl.rows[0].cells]
        if len(hdr) >= 2 and hdr[0].lower().startswith('test') and 'requirement' in hdr[1].lower():
            target = tbl; break
    if target:
        mapping = {'T‑001':'T001','T‑002':'T002','T‑003':'T003','T‑004':'T004','T-001':'T001','T-002':'T002','T-003':'T003','T-004':'T004'}
        for r in range(1, min(5, len(target.rows))):
            tid = target.cell(r,0).text.strip()
            key = mapping.get(tid)
            if not key: continue
            m = metrics.get(key, {})
            if key=='T001': obs = f"min {m.get('min'):.3f}, mean {m.get('mean'):.3f}, max {m.get('max'):.3f}" if m else ''
            elif key=='T002': obs = f"ratio≈{m.get('mean_ratio'):.3f} (mm {m.get('mean_mm'):.2f} / base {m.get('mean_baseline_mm'):.2f})" if m else ''
            elif key=='T003': obs = f"max {m.get('max_Hz'):.3f} Hz, mean {m.get('mean_Hz'):.3f} Hz" if m else ''
            elif key=='T004': obs = f"min {m.get('min_kW'):.2f}, mean {m.get('mean_kW'):.2f}, max {m.get('max_kW'):.2f} kW" if m else ''
            else: obs = ''
            status = 'PASS' if (m and m.get('pass_rate',0)>=1.0) else (f"{int(m.get('pass_rate',0)*100)}% PASS" if m else '')
            target.cell(r,2).text = obs
            target.cell(r,3).text = status
    doc.add_heading('Auto-Inserted Evidence Paths', level=2)
    for k,v in plots.items():
        doc.add_paragraph(f"{k}: {v}", style='List Bullet')
    out = docx_path.with_name(docx_path.stem + '_UPDATED.docx')
    doc.save(out)
    print(f"[OK] Updated Go/No-Go: {out}")

def update_stw(xlsx_path: Path, plots):
    if not xlsx_path or not xlsx_path.exists():
        print(f"[WARN] STW not found: {xlsx_path}")
        return
    try:
        df = pd.read_excel(xlsx_path, engine='openpyxl')
    except Exception as e:
        print(f"[WARN] STW read failed: {e}")
        return
    ts = datetime.now().strftime('%Y-%m-%d %H:%M')
    df['Updated'] = ts
    df['Evidence_Sample'] = next(iter(plots.values())) if plots else ''
    out = xlsx_path.with_name(xlsx_path.stem + '_UPDATED.xlsx')
    df.to_excel(out, index=False)
    print(f"[OK] Updated STW: {out}")

def run_reorg(src: Path, dest: Path):
    script = None
    for cand in [Path('reorganize_enterprise.py'), src/'reorganize_enterprise.py']:
        if cand.exists():
            script = cand; break
    if not script:
        print('[WARN] reorganize_enterprise.py not found; skipping enterprise packaging.')
        return
    dest.mkdir(parents=True, exist_ok=True)
    cmd = [sys.executable, str(script), '--source', str(src), '--dest', str(dest), '--apply', '--copy']
    print('[RUN]', ' '.join(cmd))
    try:
        subprocess.run(cmd, check=True)
        print('[OK] Enterprise packaging complete.')
    except subprocess.CalledProcessError as e:
        print('[WARN] Enterprise packaging failed:', e)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--source','-s', default='.', help='Project root containing package files')
    ap.add_argument('--dest','-d', default='./MicroHydroV1_ENTERPRISE', help='Output enterprise root')
    args = ap.parse_args()
    src = Path(args.source).resolve()
    dest = Path(args.dest).resolve()
    print('[INFO] Source:', src)
    print('[INFO] Dest  :', dest)
    rnd = find_rnd_workbook(src)
    if not rnd:
        print('[ERROR] Could not locate RnD workbook (RnD_*.xlsx).'); sys.exit(2)
    print('[INFO] RnD workbook:', rnd)
    csvs = discover_intake_csvs(src)
    if not csvs:
        print('[WARN] No intake_raw CSVs discovered. Continuing with current workbook.')
    else:
        nrows = import_csvs_to_workbook(csvs, rnd)
        print(f'[OK] Imported {nrows} rows into Measurements.')
    df = pd.read_excel(rnd, sheet_name='Measurements', engine='openpyxl')
    metrics = compute_metrics(df)
    print('[INFO] Metrics:', json.dumps(metrics, indent=2))
    plot_dir = src / PLOT_DIR_REL
    plots = make_plots(df, plot_dir)
    gonogo = find_first(src, ['GoNoGo_Packet_Draft.docx','Go_NoGo_Packet.docx','GoNoGo_Packet.docx'])
    if gonogo:
        update_gonogo(gonogo, metrics, plots)
    else:
        print('[WARN] Go/No-Go draft not found; skipping.')
    stw = find_first(src, ['STW_vNext_Matrix_Filled.xlsx','STW_vNext_Matrix.xlsx'])
    if stw:
        update_stw(stw, plots)
    else:
        print('[WARN] STW matrix not found; skipping.')
    run_reorg(src, dest)
    print('
=== PIPELINE SUMMARY ===')
    print('Workbook :', rnd)
    print('Plots    :', plots)
    print('Dest     :', dest)

if __name__ == '__main__':
    main()
