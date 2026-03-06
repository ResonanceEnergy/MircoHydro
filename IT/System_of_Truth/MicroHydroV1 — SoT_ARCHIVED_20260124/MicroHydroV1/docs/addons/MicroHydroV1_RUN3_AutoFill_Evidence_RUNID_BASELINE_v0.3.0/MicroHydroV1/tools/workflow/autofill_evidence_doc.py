#!/usr/bin/env python3
"""MicroHydroV1 - Auto-fill RUN-3 Evidence Doc (v0.3.0+) with Run_ID matching + baseline row detection.

What this version adds
1) Run_ID matching:
   - Filters workbook Measurements rows to only those belonging to the requested --run-id.
   - Uses Run_ID column if present; otherwise matches --run-id as a substring in Conditions or Notes.
2) Baseline row detection for T-002:
   - If Baseline_Value exists, it is used.
   - Additionally, the script attempts to locate an explicit baseline row in the workbook (same run) where
     Sensor/Conditions/Notes indicate "baseline".
   - If found, that baseline row becomes the baseline source (more traceable than a copied baseline).

Inputs
- T-002 timeseries CSV (Time_s, Pressure_Pa)
- RnD workbook (default: automation/MicroHydroV1_RnD_Export.xlsx)

Acceptance rules (fallback if Pass is blank)
- T-001: coherence >= 0.90
- T-002: ripple <= 0.70 * baseline
- T-003: drift <= 0.10 Hz
- T-004: 10 <= power <= 12 kW

Usage
  python tools/workflow/autofill_evidence_doc.py     --root .     --run-id 2026-01-22_RUN3     --t002 tests/raw/RUN3/T002_TankRipple_timeseries.csv     --workbook automation/MicroHydroV1_RnD_Export.xlsx     --out-doc docs/run3/RUN3_Evidence_2026-01-22_RUN3.docx     --out-plot tests/results/summary/T002_PSD_Peakiness.png

Optional knobs
- --run-id-column Run_ID        (if you add a Run_ID column later)
- --run-id-fields Conditions Notes   (default substring match fields)

"""

from __future__ import annotations

import argparse
import datetime as dt
import re
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple

import numpy as np
import pandas as pd
from docx import Document

BAND_LOW = 0.2
BAND_HIGH = 20.0


def compute_psd(x: np.ndarray, fs: float):
    x = x - np.nanmean(x)
    try:
        from scipy.signal import welch
        f, pxx = welch(x, fs=fs, nperseg=min(2048, len(x)))
    except Exception:
        n = len(x)
        win = np.hanning(n)
        xf = np.fft.rfft(x * win)
        pxx = (np.abs(xf) ** 2) / (fs * np.sum(win ** 2))
        f = np.fft.rfftfreq(n, d=1.0 / fs)
    return f, pxx


def peakiness_metrics(csv_path: Path, time_col='Time_s', signal_col='Pressure_Pa'):
    df = pd.read_csv(csv_path)
    if time_col not in df.columns or signal_col not in df.columns:
        raise SystemExit(f"Missing columns. Need {time_col} and {signal_col}. Got: {list(df.columns)}")

    t = pd.to_numeric(df[time_col], errors='coerce').to_numpy()
    x = pd.to_numeric(df[signal_col], errors='coerce').to_numpy()
    m = np.isfinite(t) & np.isfinite(x)
    t, x = t[m], x[m]
    if len(t) < 10:
        raise SystemExit('Not enough samples after cleaning.')

    dtv = np.diff(t)
    dtv = dtv[np.isfinite(dtv) & (dtv > 0)]
    fs = 1.0 / float(np.median(dtv)) if len(dtv) else 100.0

    f, pxx = compute_psd(x, fs)
    band = (f >= BAND_LOW) & (f <= BAND_HIGH)
    if not np.any(band):
        raise SystemExit('No PSD bins in band. Check sampling/timebase.')

    p = pxx[band]
    peak = float(np.max(p))
    med = float(np.median(p))
    peakiness = peak / max(1e-12, med)
    f_dom = float(f[band][np.argmax(p)])

    return {
        'fs': fs,
        'peakiness': peakiness,
        'f_dom': f_dom,
        'peak_psd': peak,
        'median_psd': med,
    }


def _norm_test_id(s: Any) -> str:
    if s is None:
        return ''
    return str(s).replace('‑', '-').replace('–', '-').replace('−', '-').strip()


def _safe_float(x) -> Optional[float]:
    try:
        if x is None or x == '':
            return None
        return float(x)
    except Exception:
        return None


def _parse_ts(x) -> Optional[dt.datetime]:
    if x is None or x == '':
        return None
    if isinstance(x, dt.datetime):
        return x
    s = str(x).strip()
    try:
        return dt.datetime.fromisoformat(s)
    except Exception:
        pass
    try:
        return pd.to_datetime(s).to_pydatetime()
    except Exception:
        return None


def eval_pass_rule(test_id: str, value, baseline_value=None) -> Optional[bool]:
    tid = _norm_test_id(test_id)
    v = _safe_float(value)
    b = _safe_float(baseline_value)

    if v is None:
        return None

    if tid == 'T-001':
        return v >= 0.90
    if tid == 'T-002':
        if b is None or b <= 0:
            return None
        return v <= 0.70 * b
    if tid == 'T-003':
        return v <= 0.10
    if tid == 'T-004':
        return 10.0 <= v <= 12.0
    return None


def _truthy_pass(pas) -> Optional[bool]:
    if isinstance(pas, bool):
        return pas
    if isinstance(pas, str):
        p = pas.strip().lower()
        if p in ('true','pass','yes','1'):
            return True
        if p in ('false','fail','no','0'):
            return False
    return None


def _contains_run_id(text: str, run_id: str) -> bool:
    if not text:
        return False
    return run_id.lower() in str(text).lower()


def _is_baseline_row(sensor: str, conditions: str, notes: str) -> bool:
    blob = ' '.join([str(sensor or ''), str(conditions or ''), str(notes or '')]).lower()
    return 'baseline' in blob or 'base line' in blob


def read_measurements_rows(workbook_path: Path) -> Tuple[List[Dict[str, Any]], Dict[str, int]]:
    import openpyxl
    wb = openpyxl.load_workbook(workbook_path, data_only=True)
    if 'Measurements' not in wb.sheetnames:
        raise SystemExit('Workbook missing Measurements sheet: ' + str(workbook_path))
    ws = wb['Measurements']

    headers = {}
    for c in range(1, ws.max_column + 1):
        v = ws.cell(1, c).value
        if v:
            headers[str(v).strip()] = c

    def get(r, name):
        idx = headers.get(name)
        return ws.cell(r, idx).value if idx else None

    rows=[]
    for r in range(2, ws.max_row+1):
        tid = get(r,'Test_ID')
        if not tid:
            continue
        row = {
            'Timestamp': _parse_ts(get(r,'Timestamp')),
            'Test_ID': _norm_test_id(tid),
            'Sensor': get(r,'Sensor'),
            'Value': get(r,'Value'),
            'Unit': get(r,'Unit') or '',
            'Baseline_Value': get(r,'Baseline_Value'),
            'Conditions': get(r,'Conditions') or '',
            'Notes': get(r,'Notes') or '',
            'Pass': get(r,'Pass'),
            'Run_ID': get(r,'Run_ID'),
            'RunID': get(r,'RunID'),
        }
        rows.append(row)

    return rows, headers


def filter_by_run_id(rows: List[Dict[str, Any]], run_id: str, headers: Dict[str,int], run_id_column: Optional[str], run_id_fields: List[str]) -> List[Dict[str, Any]]:
    # If an explicit Run_ID column exists and was provided, prefer it.
    if run_id_column and run_id_column in headers:
        out=[]
        for r in rows:
            rid = r.get(run_id_column)
            if rid and str(rid).strip().lower() == run_id.lower():
                out.append(r)
        return out

    # Otherwise, match run_id substring in selected fields (default: Conditions + Notes)
    out=[]
    for r in rows:
        for f in run_id_fields:
            if _contains_run_id(r.get(f,''), run_id):
                out.append(r)
                break
    return out


def select_latest(rows: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    # latest by Timestamp
    rows2=[r for r in rows if r.get('Timestamp')]
    if rows2:
        return sorted(rows2, key=lambda x: x['Timestamp'])[-1]
    return rows[-1] if rows else None


def pick_records(rows: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    """Pick best records for each test.

    For T-001/T-003/T-004: latest row.
    For T-002: select a non-baseline ripple row + a baseline row.
    """
    by_tid={}
    for tid in ('T-001','T-003','T-004'):
        rr=[r for r in rows if r['Test_ID']==tid]
        rec=select_latest(rr)
        if rec:
            by_tid[tid]=rec

    # T-002
    rr2=[r for r in rows if r['Test_ID']=='T-002']
    if rr2:
        baselines=[r for r in rr2 if _is_baseline_row(r.get('Sensor',''), r.get('Conditions',''), r.get('Notes',''))]
        non_base=[r for r in rr2 if not _is_baseline_row(r.get('Sensor',''), r.get('Conditions',''), r.get('Notes',''))]

        base=select_latest(baselines) if baselines else None
        ripple=select_latest(non_base) if non_base else select_latest(rr2)

        if ripple:
            ripple = dict(ripple)
            # determine baseline priority: explicit baseline row > Baseline_Value > None
            b_val = None
            if base and _safe_float(base.get('Value')) is not None:
                b_val = base.get('Value')
            elif _safe_float(ripple.get('Baseline_Value')) is not None:
                b_val = ripple.get('Baseline_Value')

            ripple['Baseline_Value'] = b_val
            ripple['Baseline_Row_Timestamp'] = base.get('Timestamp') if base else None
            by_tid['T-002'] = ripple

    return by_tid


def format_value(tid: str, value: Any, unit: str) -> str:
    v = _safe_float(value)
    if v is None:
        s = str(value or '').strip() or '____'
    else:
        if tid in ('T-001','T-003'):
            s = f"{v:.3f}"
        else:
            s = f"{v:.2f}"
    unit = (unit or '').strip()
    return f"{s} {unit}".strip()


def replace_in_paragraph(paragraph, mapping: dict[str, str]):
    for key, val in mapping.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(key, val)


def fill_table(doc: Document, records: Dict[str, Dict[str, Any]], peakiness: str, f_dom: str):
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        if 'Test' not in header or 'Result' not in header:
            continue

        col_test=0
        col_metric = header.index('Metric') if 'Metric' in header else 1
        col_result = header.index('Result') if 'Result' in header else 3
        col_pass = header.index('Pass/Fail') if 'Pass/Fail' in header else 4

        for row in table.rows[1:]:
            tid=_norm_test_id(row.cells[col_test].text)
            metric=row.cells[col_metric].text.strip().lower()

            # resonance row
            if tid=='T-002' and 'resonance' in metric:
                row.cells[col_result].text = peakiness
                row.cells[col_pass].text = 'n/a'
                continue

            if tid not in records:
                continue

            rec=records[tid]
            pas=_truthy_pass(rec.get('Pass'))

            if tid=='T-002' and 'ripple' in metric:
                v=_safe_float(rec.get('Value'))
                b=_safe_float(rec.get('Baseline_Value'))
                if v is not None and b is not None and b>0:
                    red=(b-v)/b*100.0
                    row.cells[col_result].text = f"{v:.2f} mm | base {b:.2f} mm | {red:.1f}%"
                    pas = pas if pas is not None else eval_pass_rule('T-002', v, b)
                    row.cells[col_pass].text = 'PASS' if pas else 'FAIL'
                else:
                    row.cells[col_result].text = format_value(tid, rec.get('Value'), rec.get('Unit'))
                    row.cells[col_pass].text = 'TBD'
                continue

            # default rows
            row.cells[col_result].text = format_value(tid, rec.get('Value'), rec.get('Unit'))
            if pas is None:
                pas = eval_pass_rule(tid, rec.get('Value'))
            row.cells[col_pass].text = 'PASS' if pas is True else 'FAIL' if pas is False else 'TBD'


def append_summary(doc: Document, run_id: str, records: Dict[str, Dict[str, Any]], peakiness: str, f_dom: str):
    doc.add_paragraph('')
    doc.add_paragraph(f'Auto-filled summary for Run_ID={run_id}:')
    for tid in ('T-001','T-002','T-003','T-004'):
        if tid in records:
            r=records[tid]
            ts=r.get('Timestamp')
            ts_s=ts.isoformat(timespec='seconds') if ts else ''
            if tid=='T-002':
                doc.add_paragraph(f"- {tid}: ripple={r.get('Value')} base={r.get('Baseline_Value')} pass={_truthy_pass(r.get('Pass'))} {ts_s}")
                bts=r.get('Baseline_Row_Timestamp')
                if bts:
                    doc.add_paragraph(f"  baseline row timestamp: {bts.isoformat(timespec='seconds')}")
            else:
                doc.add_paragraph(f"- {tid}: value={r.get('Value')} {r.get('Unit')} pass={_truthy_pass(r.get('Pass'))} {ts_s}")
    doc.add_paragraph(f"- T-002 peakiness={peakiness} dominant_hz={f_dom}")


def main(argv=None):
    ap=argparse.ArgumentParser()
    ap.add_argument('--root', default='.', help='MicroHydroV1 root')
    ap.add_argument('--run-id', required=True, help='Run ID to match (e.g., 2026-01-22_RUN3)')
    ap.add_argument('--t002', required=True, help='T-002 timeseries CSV (Time_s,Pressure_Pa)')
    ap.add_argument('--workbook', default='automation/MicroHydroV1_RnD_Export.xlsx')
    ap.add_argument('--template', default='docs/run3/RUN3_Evidence_Template.docx')
    ap.add_argument('--out-doc', required=True)
    ap.add_argument('--out-plot', default=None)

    ap.add_argument('--run-id-column', default=None, help='Optional explicit Run_ID column name')
    ap.add_argument('--run-id-fields', nargs='*', default=['Conditions','Notes'], help='Fields to search for Run_ID substring')

    args=ap.parse_args(argv)
    root=Path(args.root).resolve()

    # T-002 peakiness
    t002_path=Path(args.t002)
    if not t002_path.is_absolute():
        t002_path=root/t002_path
    t2=peakiness_metrics(t002_path)
    peakiness=f"{t2['peakiness']:.3f}"
    f_dom=f"{t2['f_dom']:.2f}"
    fs=f"{t2['fs']:.2f}"

    # Optional plot
    if args.out_plot:
        plot_script=root/'tools/plots/plot_t002_psd_peakiness.py'
        if plot_script.exists():
            import subprocess
            out_plot=Path(args.out_plot)
            if not out_plot.is_absolute():
                out_plot=root/out_plot
            out_plot.parent.mkdir(parents=True, exist_ok=True)
            subprocess.check_call(['python', str(plot_script), '--csv', str(t002_path), '--out', str(out_plot)])

    # Workbook parsing
    wb_path=Path(args.workbook)
    if not wb_path.is_absolute():
        wb_path=root/wb_path
    records={}
    matched_count=0
    if wb_path.exists():
        rows, headers = read_measurements_rows(wb_path)
        filtered = filter_by_run_id(rows, args.run_id, headers, args.run_id_column, args.run_id_fields)
        matched_count = len(filtered)
        records = pick_records(filtered)

    # Load template
    template=root/args.template
    if not template.exists():
        raise SystemExit('Template not found: '+str(template))

    doc=Document(str(template))
    now=dt.datetime.now().isoformat(timespec='seconds')
    mapping={
        '{{RUN_ID}}': args.run_id,
        '{{DATETIME}}': now,
        '{{PEAKINESS}}': peakiness,
        '{{F_DOM}}': f_dom,
        'Observed dominant peak (Hz): ____': f"Observed dominant peak (Hz): {f_dom}",
        'Sampling: 100 Hz': f"Sampling: 100 Hz (inferred fs≈{fs} Hz)",
    }
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)

    fill_table(doc, records, peakiness, f_dom)
    append_summary(doc, args.run_id, records, peakiness, f_dom)

    out_doc=Path(args.out_doc)
    if not out_doc.is_absolute():
        out_doc=root/out_doc
    out_doc.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_doc))

    print('Wrote:', out_doc)
    print('Workbook rows matched to Run_ID:', matched_count)
    print('T-002 peakiness:', peakiness, 'dominant Hz:', f_dom, 'fs≈', fs)


if __name__=='__main__':
    main()
