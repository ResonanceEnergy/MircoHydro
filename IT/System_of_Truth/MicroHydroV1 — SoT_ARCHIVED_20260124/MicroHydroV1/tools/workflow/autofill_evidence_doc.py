#!/usr/bin/env python3
"""MicroHydroV1 - Auto-fill RUN-3 Evidence Doc (v0.3.0+)

Features
- Computes T-002 resonance peakiness from time-series CSV (Time_s, Pressure_Pa)
- Optionally generates PSD plot PNG
- Pulls latest measurements from automation/MicroHydroV1_RnD_Export.xlsx (Measurements sheet):
  - T-001 coherence: value + PASS/FAIL
  - T-002 ripple reduction vs baseline: value, baseline, reduction%, PASS/FAIL
  - T-003 ELC drift: value + PASS/FAIL
  - T-004 power: value + PASS/FAIL
- Populates the Evidence Summary table in docs/run3/RUN3_Evidence_Template.docx

Notes
- openpyxl does not recalc Excel formulas. If Pass cells are blank, this script evaluates PASS/FAIL
  using acceptance rules:
    T-001: coherence >= 0.90
    T-002: ripple <= 0.70 * baseline
    T-003: drift <= 0.10 Hz
    T-004: 10 <= power <= 12 kW

Usage
  python tools/workflow/autofill_evidence_doc.py     --root .     --run-id 2026-01-22_RUN3     --t002 tests/raw/RUN3/T002_TankRipple_timeseries.csv     --workbook automation/MicroHydroV1_RnD_Export.xlsx     --out-doc docs/run3/RUN3_Evidence_2026-01-22_RUN3.docx     --out-plot tests/results/summary/T002_PSD_Peakiness.png
"""

from __future__ import annotations

import argparse
import datetime as dt
import re
from pathlib import Path
from typing import Dict, Any, Optional, Tuple

import numpy as np
import pandas as pd
from docx import Document

BAND_LOW = 0.2
BAND_HIGH = 20.0


def compute_psd(x: np.ndarray, fs: float):
    x = x - np.nanmean(x)
    try:
        from scipy.signal import welch
        f, pxx = welch(x, fs=fs, nperseg=min(2048, len(x)))
    except Exception:
        n = len(x)
        win = np.hanning(n)
        xf = np.fft.rfft(x * win)
        pxx = (np.abs(xf) ** 2) / (fs * np.sum(win ** 2))
        f = np.fft.rfftfreq(n, d=1.0 / fs)
    return f, pxx


def peakiness_metrics(csv_path: Path, time_col='Time_s', signal_col='Pressure_Pa'):
    df = pd.read_csv(csv_path)
    if time_col not in df.columns or signal_col not in df.columns:
        raise SystemExit(f"Missing columns. Need {time_col} and {signal_col}. Got: {list(df.columns)}")

    t = pd.to_numeric(df[time_col], errors='coerce').to_numpy()
    x = pd.to_numeric(df[signal_col], errors='coerce').to_numpy()
    m = np.isfinite(t) & np.isfinite(x)
    t, x = t[m], x[m]
    if len(t) < 10:
        raise SystemExit('Not enough samples after cleaning.')

    dtv = np.diff(t)
    dtv = dtv[np.isfinite(dtv) & (dtv > 0)]
    fs = 1.0 / float(np.median(dtv)) if len(dtv) else 100.0

    f, pxx = compute_psd(x, fs)
    band = (f >= BAND_LOW) & (f <= BAND_HIGH)
    if not np.any(band):
        raise SystemExit('No PSD bins in band. Check sampling/timebase.')

    p = pxx[band]
    peak = float(np.max(p))
    med = float(np.median(p))
    peakiness = peak / max(1e-12, med)
    f_dom = float(f[band][np.argmax(p)])

    return {
        'fs': fs,
        'peakiness': peakiness,
        'f_dom': f_dom,
        'peak_psd': peak,
        'median_psd': med,
    }


def _norm_test_id(s: Any) -> str:
    if s is None:
        return ''
    return str(s).replace('‑', '-').replace('–', '-').replace('−', '-').strip()


def _safe_float(x) -> Optional[float]:
    try:
        if x is None or x == '':
            return None
        return float(x)
    except Exception:
        return None


def _parse_ts(x) -> Optional[dt.datetime]:
    if x is None or x == '':
        return None
    if isinstance(x, dt.datetime):
        return x
    s = str(x).strip()
    try:
        return dt.datetime.fromisoformat(s)
    except Exception:
        pass
    try:
        return pd.to_datetime(s).to_pydatetime()
    except Exception:
        return None


def eval_pass_rule(test_id: str, sensor: str, value, unit: str, baseline_value=None) -> Optional[bool]:
    tid = _norm_test_id(test_id)
    s = (sensor or '').lower()
    v = _safe_float(value)
    b = _safe_float(baseline_value)

    if v is None:
        return None

    if tid == 'T-001':
        return v >= 0.90

    if tid == 'T-002':
        if b is None or b <= 0:
            return None
        return v <= 0.70 * b

    if tid == 'T-003':
        return v <= 0.10

    if tid == 'T-004':
        return 10.0 <= v <= 12.0

    return None


def _parse_baseline_from_text(text: str) -> Optional[float]:
    if not text:
        return None
    # patterns: Baseline=12.4 or Baseline:12.4
    m = re.search(r'Baseline\s*[:=]\s*([0-9]+(?:\.[0-9]+)?)', text, flags=re.IGNORECASE)
    if m:
        return _safe_float(m.group(1))
    return None


def read_workbook_latest(workbook_path: Path) -> Dict[str, Dict[str, Any]]:
    """Read Measurements sheet and return latest row per Test_ID.

    For T-002, we keep the latest row that has a usable baseline.
    """
    import openpyxl

    wb = openpyxl.load_workbook(workbook_path, data_only=True)
    if 'Measurements' not in wb.sheetnames:
        raise SystemExit('Workbook missing Measurements sheet: ' + str(workbook_path))
    ws = wb['Measurements']

    # map headers
    headers = {}
    for c in range(1, ws.max_column + 1):
        v = ws.cell(1, c).value
        if v:
            headers[str(v).strip()] = c

    def col(name):
        return headers.get(name)

    idx_ts = col('Timestamp')
    idx_tid = col('Test_ID')
    idx_sensor = col('Sensor')
    idx_val = col('Value')
    idx_unit = col('Unit')
    idx_pass = col('Pass')
    idx_base = col('Baseline_Value')
    idx_cond = col('Conditions')
    idx_notes = col('Notes')

    if not all([idx_ts, idx_tid, idx_sensor, idx_val]):
        raise SystemExit('Measurements sheet missing required headers.')

    best: Dict[str, Dict[str, Any]] = {}

    for r in range(2, ws.max_row + 1):
        tid = ws.cell(r, idx_tid).value
        if not tid:
            continue
        tid_n = _norm_test_id(tid)
        if tid_n not in ('T-001','T-002','T-003','T-004'):
            continue

        ts = _parse_ts(ws.cell(r, idx_ts).value)
        sensor = ws.cell(r, idx_sensor).value
        val = ws.cell(r, idx_val).value
        unit = ws.cell(r, idx_unit).value if idx_unit else ''
        pas = ws.cell(r, idx_pass).value if idx_pass else None

        baseline = ws.cell(r, idx_base).value if idx_base else None
        if baseline is None:
            # sometimes baseline is embedded in Notes or Conditions
            baseline = _parse_baseline_from_text(str(ws.cell(r, idx_notes).value or '')) if idx_notes else None
        if baseline is None:
            baseline = _parse_baseline_from_text(str(ws.cell(r, idx_cond).value or '')) if idx_cond else None

        # normalize pass
        if isinstance(pas, str):
            pnorm = pas.strip().lower()
            if pnorm in ('true','pass','yes','1'):
                pas = True
            elif pnorm in ('false','fail','no','0'):
                pas = False
            else:
                pas = None

        # if pass empty, compute
        if pas is None:
            pas = eval_pass_rule(tid_n, str(sensor or ''), val, str(unit or ''), baseline)

        rec = {
            'timestamp': ts,
            'sensor': sensor,
            'value': val,
            'unit': unit,
            'pass': pas,
            'baseline': baseline,
        }

        def is_better(existing: Dict[str, Any]) -> bool:
            old_ts = existing.get('timestamp')
            if ts and (old_ts is None or ts > old_ts):
                return True
            return False

        if tid_n == 'T-002':
            # only accept rows with baseline
            if _safe_float(baseline) is None:
                continue

        if tid_n not in best:
            best[tid_n] = rec
        else:
            if is_better(best[tid_n]):
                best[tid_n] = rec

    return best


def replace_in_paragraph(paragraph, mapping: dict[str, str]):
    for key, val in mapping.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(key, val)


def _format_value(tid: str, value: Any, unit: str) -> str:
    v = _safe_float(value)
    if v is None:
        s = str(value or '').strip() or '____'
    else:
        if tid == 'T-001':
            s = f"{v:.3f}"
        elif tid == 'T-003':
            s = f"{v:.3f}"
        else:
            s = f"{v:.2f}"
    unit = (unit or '').strip()
    return f"{s} {unit}".strip()


def fill_table_values(doc: Document, wb_metrics: Dict[str, Dict[str, Any]], peakiness: str, f_dom: str):
    # Fill the summary table
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        if 'Test' not in header or 'Result' not in header:
            continue

        col_test = 0
        col_metric = header.index('Metric') if 'Metric' in header else 1
        col_req = header.index('Requirement') if 'Requirement' in header else 2
        col_result = header.index('Result') if 'Result' in header else 3
        col_pass = header.index('Pass/Fail') if 'Pass/Fail' in header else 4

        for row in table.rows[1:]:
            tid = _norm_test_id(row.cells[col_test].text)
            metric = row.cells[col_metric].text.strip().lower()

            # T-002 resonance row
            if tid == 'T-002' and 'resonance' in metric:
                row.cells[col_result].text = peakiness
                row.cells[col_pass].text = 'n/a'
                continue

            if tid in wb_metrics:
                rec = wb_metrics[tid]

                # T-002 ripple reduction row
                if tid == 'T-002' and 'ripple' in metric:
                    v = _safe_float(rec.get('value'))
                    b = _safe_float(rec.get('baseline'))
                    if v is not None and b and b > 0:
                        red = (b - v) / b * 100.0
                        row.cells[col_result].text = f"{v:.2f} mm | base {b:.2f} mm | {red:.1f}%"
                        pas = rec.get('pass')
                        row.cells[col_pass].text = 'PASS' if pas is True else 'FAIL' if pas is False else 'TBD'
                    else:
                        row.cells[col_result].text = _format_value(tid, rec.get('value'), rec.get('unit'))
                        row.cells[col_pass].text = 'TBD'
                    continue

                # default fill for T-001/T-003/T-004
                row.cells[col_result].text = _format_value(tid, rec.get('value'), rec.get('unit'))
                pas = rec.get('pass')
                row.cells[col_pass].text = 'PASS' if pas is True else 'FAIL' if pas is False else 'TBD'

    # Append a summary block
    doc.add_paragraph('')
    doc.add_paragraph('Auto-filled summary:')
    for tid in ('T-001','T-002','T-003','T-004'):
        if tid in wb_metrics:
            rec = wb_metrics[tid]
            ts = rec.get('timestamp')
            ts_s = ts.isoformat(timespec='seconds') if ts else ''
            v = rec.get('value'); unit = rec.get('unit') or ''
            pas = rec.get('pass')
            if tid == 'T-002':
                b = rec.get('baseline')
                doc.add_paragraph(f"- {tid}: value={v} {unit} baseline={b} PASS={pas} {ts_s}")
            else:
                doc.add_paragraph(f"- {tid}: value={v} {unit} PASS={pas} {ts_s}")
    doc.add_paragraph(f"- T-002 peakiness: {peakiness} | dominant Hz: {f_dom}")


def main(argv=None):
    ap = argparse.ArgumentParser()
    ap.add_argument('--root', default='.', help='MicroHydroV1 root')
    ap.add_argument('--run-id', required=True)
    ap.add_argument('--t002', required=True, help='Path to T-002 timeseries CSV')
    ap.add_argument('--out-doc', required=True, help='Output populated docx')
    ap.add_argument('--template', default='docs/run3/RUN3_Evidence_Template.docx')
    ap.add_argument('--out-plot', default=None, help='Optional: output plot PNG by calling plot script')
    ap.add_argument('--workbook', default='automation/MicroHydroV1_RnD_Export.xlsx', help='RnD workbook path')
    args = ap.parse_args(argv)

    root = Path(args.root).resolve()

    # T-002 peakiness
    t002_path = Path(args.t002)
    if not t002_path.is_absolute():
        t002_path = root / t002_path

    t2 = peakiness_metrics(t002_path)
    peakiness = f"{t2['peakiness']:.3f}"
    f_dom = f"{t2['f_dom']:.2f}"
    fs = f"{t2['fs']:.2f}"

    # optional plot
    if args.out_plot:
        plot_script = root/'tools/plots/plot_t002_psd_peakiness.py'
        if plot_script.exists():
            import subprocess
            out_plot = Path(args.out_plot)
            if not out_plot.is_absolute():
                out_plot = root/out_plot
            out_plot.parent.mkdir(parents=True, exist_ok=True)
            subprocess.check_call(['python', str(plot_script), '--csv', str(t002_path), '--out', str(out_plot)])

    # workbook metrics
    wb_path = Path(args.workbook)
    if not wb_path.is_absolute():
        wb_path = root / wb_path

    wb_metrics: Dict[str, Dict[str, Any]] = {}
    if wb_path.exists():
        wb_metrics = read_workbook_latest(wb_path)

    # load template
    template = root / args.template
    if not template.exists():
        raise SystemExit('Template not found: ' + str(template))

    doc = Document(str(template))
    now = dt.datetime.now().isoformat(timespec='seconds')

    mapping = {
        '{{RUN_ID}}': args.run_id,
        '{{DATETIME}}': now,
        '{{PEAKINESS}}': peakiness,
        '{{F_DOM}}': f_dom,
        'Observed dominant peak (Hz): ____': f"Observed dominant peak (Hz): {f_dom}",
        'Sampling: 100 Hz': f"Sampling: 100 Hz (inferred fs≈{fs} Hz)",
    }

    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)

    fill_table_values(doc, wb_metrics=wb_metrics, peakiness=peakiness, f_dom=f_dom)

    out_doc = Path(args.out_doc)
    if not out_doc.is_absolute():
        out_doc = root/out_doc
    out_doc.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_doc))

    print('Wrote:', out_doc)
    print('T-002 peakiness:', peakiness, 'dominant Hz:', f_dom, 'fs≈', fs)
    for tid, rec in wb_metrics.items():
        print('Workbook', tid, 'value=', rec.get('value'), 'baseline=', rec.get('baseline'), 'pass=', rec.get('pass'))


if __name__ == '__main__':
    main()
